"""강의안 마크다운 → DOCX 변환 스크립트"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 페이지 설정 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 스타일 설정 ──
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '맑은 고딕'
    hs._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    hs.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    elif level == 3:
        hs.font.size = Pt(13)
    else:
        hs.font.size = Pt(11)

NAVY = RGBColor(0x1a, 0x23, 0x7e)
ORANGE = RGBColor(0xff, 0x8f, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)

# ── 마크다운 파싱 ──
with open('강의안_바이브코딩_3시간.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

def add_styled_text(paragraph, text):
    """볼드/코드 인라인 마크다운 처리"""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = NAVY
        else:
            paragraph.add_run(part)

def add_table_from_lines(table_lines):
    """마크다운 테이블 → DOCX 테이블"""
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells and not all(set(c) <= set('- :') for c in cells):
            rows_data.append(cells)
    if not rows_data:
        return
    cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                if i == 0:
                    run.bold = True
                    shading = cell._element.get_or_add_tcPr()
                    bg = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): '1a237e',
                        qn('w:val'): 'clear'
                    })
                    shading.append(bg)
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

in_code_block = False
code_lines = []
in_table = False
table_lines = []
i = 0

while i < len(lines):
    line = lines[i]
    stripped = line.rstrip('\n')

    # 코드 블록
    if stripped.startswith('```'):
        if in_code_block:
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 배경색 효과
            shading = p._element.get_or_add_pPr()
            bg = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'F5F5F5',
                qn('w:val'): 'clear'
            })
            shading.append(bg)
            in_code_block = False
            code_lines = []
        else:
            if in_table:
                add_table_from_lines(table_lines)
                table_lines = []
                in_table = False
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_lines.append(stripped)
        i += 1
        continue

    # 테이블
    if '|' in stripped and stripped.strip().startswith('|'):
        if not in_table:
            in_table = True
        table_lines.append(stripped)
        i += 1
        continue
    elif in_table:
        add_table_from_lines(table_lines)
        table_lines = []
        in_table = False

    # 빈 줄
    if not stripped.strip():
        i += 1
        continue

    # 구분선
    if stripped.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        # 하단 테두리
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '6',
            qn('w:space'): '1',
            qn('w:color'): 'CCCCCC'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    # 헤딩
    heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        h = doc.add_heading(level=level)
        add_styled_text(h, text)
        i += 1
        continue

    # 인용(블록쿼트)
    if stripped.startswith('>'):
        text = re.sub(r'^>\s*', '', stripped)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # 왼쪽 테두리
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        left = pBdr.makeelement(qn('w:left'), {
            qn('w:val'): 'single',
            qn('w:sz'): '18',
            qn('w:space'): '8',
            qn('w:color'): 'FF8F00'
        })
        pBdr.append(left)
        pPr.append(pBdr)
        add_styled_text(p, text)
        for run in p.runs:
            run.font.color.rgb = GRAY
            run.font.size = Pt(9.5)
        i += 1
        continue

    # 체크리스트
    if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
        checked = stripped.startswith('- [x]')
        text = stripped[5:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        marker = '☑ ' if checked else '☐ '
        run = p.add_run(marker)
        run.font.size = Pt(10)
        add_styled_text(p, text)
        i += 1
        continue

    # 리스트 (번호/불릿)
    list_match = re.match(r'^(\s*)(\d+\.|[-*])\s+(.*)', stripped)
    if list_match:
        indent = len(list_match.group(1))
        marker = list_match.group(2)
        text = list_match.group(3)
        p = doc.add_paragraph()
        indent_cm = 0.8 + (indent // 2) * 0.5
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if marker in ['-', '*']:
            p.add_run('• ')
        else:
            p.add_run(f'{marker} ')
        add_styled_text(p, text)
        i += 1
        continue

    # 일반 텍스트
    p = doc.add_paragraph()
    add_styled_text(p, stripped)
    i += 1

# 잔여 테이블
if in_table:
    add_table_from_lines(table_lines)

# ── 저장 ──
output = '강의안_바이브코딩_3시간.docx'
doc.save(output)
print(f'DOCX 생성 완료: {output}')
